import os
from pathlib import Path
from PIL import Image
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
from pdf2image import convert_from_path
import docx
from docx.shared import Pt
import threading
import asyncio

# Lazy-loaded globals to save memory (especially important for Render deploy)
_processor = None
_model = None
_model_lock = threading.Lock()

def get_trocr_model():
    """Lazy loads the HuggingFace TrOCR model (microsoft/trocr-base-printed)."""
    global _processor, _model
    if _model is None or _processor is None:
        with _model_lock:
            # Double check inside lock
            if _model is None or _processor is None:
                print("Loading TrOCR model into memory...")
                model_name = "microsoft/trocr-base-printed"
                _processor = TrOCRProcessor.from_pretrained(model_name)
                _model = VisionEncoderDecoderModel.from_pretrained(model_name)
                print("TrOCR model loaded.")
    return _processor, _model

def extract_text_from_image(image: Image.Image) -> str:
    processor, model = get_trocr_model()
    
    # Preprocess image
    image.thumbnail((1200, 1600), Image.Resampling.LANCZOS)
    if image.mode != "RGB":
        image = image.convert("RGB")
        
    pixel_values = processor(images=image, return_tensors="pt").pixel_values
    
    # Generate text
    generated_ids = model.generate(pixel_values, max_new_tokens=500)
    generated_text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
    
    return generated_text

async def convert_scanned_pdf_to_docx_ocr(input_pdf: Path, output_dir: Path) -> Path:
    """
    Extracts text from scanned PDF using OCR and saves as DOCX.
    Processes pages sequentially to prevent excessive memory usage.
    """
    try:
        loop = asyncio.get_running_loop()
        
        # 1. Convert PDF to images in thread pool (dpi=150 saves memory vs 300)
        images = await loop.run_in_executor(None, lambda: convert_from_path(str(input_pdf), dpi=150))
        
        # 2. Extract text (sequentially to control memory)
        extracted_pages = []
        for img in images:
            # Evaluate each page on a separate thread to unblock the async loop
            page_text = await loop.run_in_executor(None, extract_text_from_image, img)
            extracted_pages.append(page_text)
            
        # 3. Create DOCX
        doc = docx.Document()
        for i, text in enumerate(extracted_pages):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
            if i < len(extracted_pages) - 1:
                doc.add_page_break()
                
        output_path = output_dir / f"{input_pdf.stem}.docx"
        doc.save(str(output_path))
        
        return output_path
        
    except Exception as e:
        raise RuntimeError(f"OCR processing failed: {e}")
